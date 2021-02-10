import os
import httplib2
from docxtpl import DocxTemplate

import apiclient.discovery
from oauth2client.service_account import ServiceAccountCredentials

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


sheet = '2017'
spreadsheetId  = '1SEz5_QVrWDzSb6g60KlwxShdmPGHSnq3EBw5zJ8k-jQ'

CREDENTIALS_FILE = 'annotation2-c46cdd2a61d8.json'


def create_table(dict_:dict, name_docx):
    '''This function adds a tag to the table[0] from template.docx'''

    document = Document(f'template.docx')
    long_topics = int((len(dict_)-19)/3)
    table = document.tables[0]
    for i in range(long_topics):
        new_row = table.add_row()
            # first column 
        paragraph = table.rows[i+1].cells[0].add_paragraph()
        first_column = paragraph.add_run(f'{str(i+1)}.')
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first_column.bold = True
        first_column.font.name = 'Calibri'
        first_column.font.size = Pt(10)      

            # second_column
        cell = table.rows[i+1].cells[1]
        cell.text = f'{{{{topic{i+1}}}}}'   
        format_cell = cell.paragraphs[0].runs[0]
        format_cell.font.name = 'Calibri'
        format_cell.font.size = Pt(10)
            
            # third column
        cell = table.rows[i+1].cells[2]
        cell.text = f'{{{{content_topic{i+1}}}}}' 
        format_cell = cell.paragraphs[0].runs[0]  
        format_cell.font.name = 'Calibri'
        format_cell.font.size = Pt(10)

            # fourth_column
        cell = table.rows[i+1].cells[3]
        cell.text = f'{{{{competence{i+1}}}}}' 
        format_cell = cell.paragraphs[0].runs[0]  
        format_cell.font.name = 'Calibri'
        format_cell.font.size = Pt(10)

    document.save(f'annotations/{name_docx}.docx')       

def generate_docx_from_schedule(context:dict, lenght:int, count:int):
    '''This function changes the tag to a word from the google sheet table'''
    try:
        os.mkdir('annotations')
        print('Folder is create!')
    except OSError as e:
        temp = context['name_discipline']
        index = context['index']
        name_lenght = temp.find('/')
        name_docx = str(index+'_'+temp[:name_lenght])
        try:
            create_table(context, name_docx)
            template = DocxTemplate(f'annotations/{name_docx}.docx')
            template.render(context)
            template.save("annotations/%s.docx" %str(index+'_'+temp[:name_lenght]))
        except:
            print("Oops!")
        return print(f'successful create {count}/{lenght}')



if __name__=='__main__':       
    ## connect API google sheets
    credentials = ServiceAccountCredentials.from_json_keyfile_name(CREDENTIALS_FILE, 
                    ['https://www.googleapis.com/auth/spreadsheets',
                    'https://www.googleapis.com/auth/drive'])

    httpAuth = credentials.authorize(httplib2.Http()) 
    service = apiclient.discovery.build('sheets', 'v4', http = httpAuth)
    range_name = f'{sheet}!A2:AF3'

    #filter list of sheet
    sheet = service.spreadsheets().values().get(spreadsheetId=spreadsheetId, range=range_name).execute().get('values', [])
    for num in range(len(sheet)):
        if '0' in sheet[num][0]:
            sheet.pop(num)

    # clear cells 
    data_sheet = []
    for row in range(len(sheet)):
        temp = []
        for cell in range(len(sheet[row])):
            new_cell = sheet[row][cell].replace(sheet[row][cell],sheet[row][cell].strip())  
            temp.append(new_cell)
            if cell == len(sheet[row])-1:
                data_sheet.append(temp)


    #future keys
    keys = ['filter','index','name_discipline','description','block', 'code','stream',
                    'course_och', 'semester_och', 'course_z', 'form_educational', 'credit_hours', 
                'academic_hours', 'countact_hours', 'credit_hours_z', 'academic_hours_z',
                    'countact_hours_z', 'topic1', 'content_topic1', 'competence1', 
                    'topic2', 'content_topic2', 'competence2', 'topic3', 'content_topic3', 'competence3',
                    'topic4', 'content_topic4', 'competence4', 'topic5', 'content_topic5', 'competence5',
                    'topic6', 'content_topic6', 'competence6']

    # match two list
    dict_ = {}
    for i, value in enumerate(data_sheet):
        new_match_list = dict(zip(keys, value))              
        dict_.setdefault(i, new_match_list)

    #clear ram 
    sheet.clear()
    temp.clear()
    data_sheet.clear()

    ## add new parametors 
    for index in dict_:
        if 'В.ДВ.' in dict_[index]['index']:       
            choice_type = {'choice_type':'части, формируемой участниками образовательных отношений'} 
            type_discipline = {'type_discipline':'по выбору'} 
            dict_[index].update(choice_type)
            dict_[index].update(type_discipline)
        else:
            choice_type = {'choice_type':'обязательной части'} 
            type_discipline = {'type_discipline':'обязательной'} 
            dict_[index].update(choice_type)
            dict_[index].update(type_discipline)

    #generate process
    for i in range(0,len(dict_)):
        generate_docx_from_schedule(dict_[i], len(dict_), i+1)